from docx import Document
import os, sys

def clean_document_by_product(master_path, product_heading, output_path):
    doc = Document(master_path)
    new_doc = Document()
    new_doc._body.clear_content()

    copying = False

    paragraphs = iter(doc.paragraphs)
    tables = iter(doc.tables)
    table_elements = [t._element for t in doc.tables]
    para_elements = [p._element for p in doc.paragraphs]

    for block in doc.element.body:
        if block in para_elements:
            para = next(paragraphs)
            if para.style.name == "Heading 1":
                copying = product_heading.lower() in para.text.lower()

            if copying: # start copying when the product heading is found
                new_para = new_doc.add_paragraph(style=para.style)
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline

        elif block in table_elements:
            table = next(tables) # copies the table if copying is True
            if copying:
                new_table = new_doc.add_table(rows=0, cols=len(table.columns))
                new_table.style = 'Table Grid'
                for row in table.rows:
                    row_cells = new_table.add_row().cells
                    for i, cell in enumerate(row.cells):
                        row_cells[i].text = cell.text

    new_doc.save(output_path) # save the cleaned document
    print(f"Document saved for product '{product_heading}' at {output_path}")

if __name__ == "__main__":
    if len(sys.argv) != 2:
        sys.exit(1)

    selected_product = sys.argv[1]
    master_doc = "master.docx"
    output_file = os.path.join("output", f"{selected_product.lower().replace(' ', '_')}_cleaned.docx")
    os.makedirs("output", exist_ok=True)

    clean_document_by_product(master_doc, selected_product, output_file)